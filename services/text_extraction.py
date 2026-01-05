"""Text extraction service for various file types."""
import os
from pathlib import Path
from typing import Optional, Dict, Any
import mimetypes

# PDF
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# DOCX
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# MSG
try:
    import extract_msg
except ImportError:
    extract_msg = None

# Email
try:
    import email
    from email import policy
except ImportError:
    email = None

# CSV
try:
    import pandas as pd
except ImportError:
    pd = None

# Images/OCR
try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

from config import settings


class TextExtractionService:
    """Service for extracting text from various file types."""
    
    def __init__(self):
        self.enable_ocr = settings.enable_ocr
        self.ocr_language = settings.ocr_language
    
    def extract_text(self, file_path: Path, mime_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract text from a file.
        
        Returns:
            Dict with keys: raw_text, extracted_text, metadata
        """
        if mime_type is None:
            mime_type, _ = mimetypes.guess_type(str(file_path))
        
        # Determine file type from extension if mime_type is unknown
        if not mime_type:
            ext = file_path.suffix.lower()
            mime_type = self._get_mime_from_extension(ext)
        
        result = {
            'raw_text': '',
            'extracted_text': '',
            'metadata': {}
        }
        
        try:
            if mime_type == 'application/pdf':
                result = self._extract_pdf(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                result = self._extract_docx(file_path)
            elif mime_type == 'application/vnd.ms-outlook' or file_path.suffix.lower() == '.msg':
                result = self._extract_msg(file_path)
            elif mime_type == 'message/rfc822' or file_path.suffix.lower() == '.eml':
                result = self._extract_eml(file_path)
            elif mime_type == 'text/plain' or file_path.suffix.lower() == '.txt':
                result = self._extract_txt(file_path)
            elif mime_type == 'text/csv' or file_path.suffix.lower() == '.csv':
                result = self._extract_csv(file_path)
            elif mime_type and mime_type.startswith('image/'):
                result = self._extract_image(file_path)
            else:
                # Try as text file
                result = self._extract_txt(file_path)
        
        except Exception as e:
            result['metadata']['extraction_error'] = str(e)
            result['raw_text'] = f"Error extracting text: {str(e)}"
        
        # Clean extracted text
        if result['raw_text']:
            result['extracted_text'] = self._clean_text(result['raw_text'])
        
        return result
    
    def _extract_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 not installed")
        
        text_parts = []
        metadata = {}
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            metadata = {
                'num_pages': len(pdf_reader.pages),
                'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                'subject': pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else '',
            }
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    text_parts.append(text)
                except Exception as e:
                    text_parts.append(f"[Page {page_num + 1} extraction error: {str(e)}]")
        
        raw_text = '\n\n'.join(text_parts)
        return {
            'raw_text': raw_text,
            'extracted_text': '',
            'metadata': metadata
        }
    
    def _extract_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX."""
        if DocxDocument is None:
            raise ImportError("python-docx not installed")
        
        doc = DocxDocument(file_path)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                text_parts.append(row_text)
        
        metadata = {
            'core_properties': {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
            }
        }
        
        raw_text = '\n'.join(text_parts)
        return {
            'raw_text': raw_text,
            'extracted_text': '',
            'metadata': metadata
        }
    
    def _extract_msg(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from MSG (Outlook message)."""
        if extract_msg is None:
            raise ImportError("extract-msg not installed")
        
        msg = extract_msg.Message(str(file_path))
        
        text_parts = []
        if msg.body:
            text_parts.append(msg.body)
        if msg.htmlBody:
            text_parts.append(msg.htmlBody)
        
        metadata = {
            'sender': msg.sender,
            'to': msg.to,
            'cc': msg.cc,
            'bcc': msg.bcc,
            'subject': msg.subject,
            'date': msg.date.isoformat() if msg.date else None,
            'attachments': [att.longFilename for att in msg.attachments] if msg.attachments else [],
        }
        
        raw_text = '\n\n'.join(text_parts)
        return {
            'raw_text': raw_text,
            'extracted_text': '',
            'metadata': metadata
        }
    
    def _extract_eml(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from EML (email file)."""
        if email is None:
            raise ImportError("email parser not available")
        
        with open(file_path, 'rb') as f:
            msg = email.message_from_bytes(f.read(), policy=policy.default)
        
        text_parts = []
        html_parts = []
        
        # Extract body
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == 'text/plain':
                    text_parts.append(part.get_payload(decode=True).decode('utf-8', errors='ignore'))
                elif content_type == 'text/html':
                    html_parts.append(part.get_payload(decode=True).decode('utf-8', errors='ignore'))
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                text_parts.append(payload.decode('utf-8', errors='ignore'))
        
        metadata = {
            'from': msg.get('From', ''),
            'to': msg.get('To', ''),
            'cc': msg.get('Cc', ''),
            'bcc': msg.get('Bcc', ''),
            'subject': msg.get('Subject', ''),
            'date': msg.get('Date', ''),
        }
        
        raw_text = '\n\n'.join(text_parts + html_parts)
        return {
            'raw_text': raw_text,
            'extracted_text': '',
            'metadata': metadata
        }
    
    def _extract_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text file."""
        import chardet
        
        with open(file_path, 'rb') as f:
            raw_bytes = f.read()
        
        # Detect encoding
        detected = chardet.detect(raw_bytes)
        encoding = detected.get('encoding', 'utf-8')
        
        try:
            raw_text = raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            raw_text = raw_bytes.decode('utf-8', errors='ignore')
        
        return {
            'raw_text': raw_text,
            'extracted_text': '',
            'metadata': {'encoding': encoding}
        }
    
    def _extract_csv(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from CSV file."""
        if pd is None:
            raise ImportError("pandas not installed")
        
        try:
            df = pd.read_csv(file_path)
            # Convert to text representation
            raw_text = df.to_string(index=False)
            metadata = {
                'num_rows': len(df),
                'columns': df.columns.tolist(),
            }
        except Exception as e:
            # Fallback to plain text extraction
            return self._extract_txt(file_path)
        
        return {
            'raw_text': raw_text,
            'extracted_text': '',
            'metadata': metadata
        }
    
    def _extract_image(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from image using OCR."""
        if not self.enable_ocr:
            return {
                'raw_text': '',
                'extracted_text': '',
                'metadata': {'ocr_disabled': True}
            }
        
        if Image is None or pytesseract is None:
            return {
                'raw_text': '',
                'extracted_text': '',
                'metadata': {'ocr_not_available': True}
            }
        
        try:
            image = Image.open(file_path)
            raw_text = pytesseract.image_to_string(image, lang=self.ocr_language)
            metadata = {
                'ocr_used': True,
                'image_size': image.size,
                'image_format': image.format,
            }
        except Exception as e:
            raw_text = ''
            metadata = {'ocr_error': str(e)}
        
        return {
            'raw_text': raw_text,
            'extracted_text': '',
            'metadata': metadata
        }
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ''
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            cleaned_line = ' '.join(line.split())
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
        
        return '\n'.join(cleaned_lines)
    
    def _get_mime_from_extension(self, ext: str) -> str:
        """Get MIME type from file extension."""
        mime_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.msg': 'application/vnd.ms-outlook',
            '.eml': 'message/rfc822',
            '.txt': 'text/plain',
            '.csv': 'text/csv',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.tiff': 'image/tiff',
            '.bmp': 'image/bmp',
        }
        return mime_map.get(ext.lower(), 'application/octet-stream')

